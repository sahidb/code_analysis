# pip install radon and pylint

import os
import statistics
import argparse
from radon.complexity import cc_visit
from radon.metrics import h_visit
from radon.raw import analyze
from pylint.lint import Run
from datetime import datetime
from docx import Document
from docx.shared import RGBColor

def analyze_code(file_path):
    # Read the content of the code
    with open(file_path, 'r', encoding='utf-8') as f:
        code_content = f.read()
        
    # Cyclomatic Complexity
    cyclomatic_complexity = cc_visit(code_content)
    cc_scores = [block.complexity for block in cyclomatic_complexity]
    avg_cc = statistics.mean(cc_scores) if cc_scores else 0
    max_cc = max(cc_scores) if cc_scores else 0
    
    # Halstead Metrics
    halstead_metrics_list = list(h_visit(code_content))
    
    if halstead_metrics_list:
        halstead_metrics = halstead_metrics_list[0]
        total_operators = halstead_metrics.h1
        total_operands = halstead_metrics.h2
        vocabulary = halstead_metrics.vocabulary
        length = halstead_metrics.length
        volume = halstead_metrics.volume
        difficulty = halstead_metrics.difficulty
        effort = halstead_metrics.effort
    else:
        total_operators = total_operands = vocabulary = length = volume = difficulty = effort = 0

    # Maintainability Index
    maintainability_index = 171 - 5.2 * len(code_content.splitlines()) - 0.23 * volume - 16.2 * avg_cc if cc_scores else 100
    maintainability_index = max(0, min(maintainability_index, 100))
    
    # PEP-8 Compliance using Pylint
    try:
        pylint_results = Run([file_path], exit=False)
        pylint_score = pylint_results.linter.stats.global_note
    except Exception as e:
        print(f"Error running pylint on {file_path}: {e}")
        pylint_score = 0  # Default to 0 if pylint fails
    
    # Raw Metrics
    raw_metrics = analyze(code_content)
    
    return {
        "file_path": file_path,
        "Cyclomatic Complexity (avg)": avg_cc,
        "Cyclomatic Complexity (max)": max_cc,
        "Halstead Vocabulary": vocabulary,
        "Halstead Length": length,
        "Halstead Volume": volume,
        "Halstead Effort": effort,
        "Halstead Difficulty": difficulty,
        "Total Operators": total_operators,
        "Total Operands": total_operands,
        "Lines of Code": raw_metrics.loc,
        "Comments": raw_metrics.comments,
        "Maintainability Index": maintainability_index,
        "PEP-8 Score": pylint_score
    }

def analyze_project_folder(folder_path, exclude_folders=None):
    if exclude_folders is None:
        exclude_folders = ['venv']  # Default folder to exclude

    all_metrics = []
    
    for root, dirs, files in os.walk(folder_path):
        dirs[:] = [d for d in dirs if d not in exclude_folders]
        
        for file in files:
            if file.endswith(".py"):
                file_path = os.path.join(root, file)
                print(f"Analyzing: {file_path}")
                metrics = analyze_code(file_path)
                all_metrics.append(metrics)
    
    if not all_metrics:
        print("No Python files found in the specified folder.")
        
    return all_metrics

def aggregate_metrics(all_metrics):
    if not all_metrics:
        print("No metrics to aggregate - no Python files were analyzed.")
        return {}
    
    aggregated = {
        "Total Files": len(all_metrics),
        "Average Cyclomatic Complexity": statistics.mean([m["Cyclomatic Complexity (avg)"] for m in all_metrics]),
        "Maximum Cyclomatic Complexity": max([m["Cyclomatic Complexity (max)"] for m in all_metrics]),
        "Average Halstead Volume": statistics.mean([m["Halstead Volume"] for m in all_metrics]),
        "Average Halstead Effort": statistics.mean([m["Halstead Effort"] for m in all_metrics]),
        "Average Halstead Difficulty": statistics.mean([m["Halstead Difficulty"] for m in all_metrics]),
        "Total Lines of Code": sum([m["Lines of Code"] for m in all_metrics]),
        "Total Comments": sum([m["Comments"] for m in all_metrics]),
        "Average Maintainability Index": statistics.mean([m["Maintainability Index"] for m in all_metrics]),
        "Average PEP-8 Score": statistics.mean([m["PEP-8 Score"] for m in all_metrics])
    }
    return aggregated

def color_text(paragraph, text, color):
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(*color)

def generate_word_report(all_file_metrics, aggregated_metrics, output_path):
    if not aggregated_metrics:
        print("No data to generate a Word report.")
        return
    
    doc = Document()
    doc.add_heading('Code Analysis Report', 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    explanation = (
        "This report provides an analysis of the Python code within the specified folder. "
        "The analysis includes various metrics that give insights into the code's quality, maintainability, and complexity. "
        "Below is a brief explanation of the key metrics included in this report:\n\n"
        "- **Cyclomatic Complexity**: Measures the complexity of the code.\n"
        "- **Maintainability Index**: Indicates how maintainable the code is.\n"
        "- **PEP-8 Score**: Represents adherence to Python's style guide.\n"
        "- **Halstead Metrics**: Provides insights into code volume, effort, and difficulty."
    )
    doc.add_paragraph(explanation)

    doc.add_heading('Aggregated Project Metrics', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'

    for metric, value in aggregated_metrics.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(metric)
        row_cells[1].text = f"{value:.2f}" if isinstance(value, float) else str(value)
    
    doc.add_heading('Individual File Metrics', level=1)
    for file_metrics in all_file_metrics:
        doc.add_heading(f"File: {file_metrics['file_path']}", level=2)
        summary_paragraph = doc.add_paragraph("Summary: ")
        
        if file_metrics["Maintainability Index"] >= 80:
            color_text(summary_paragraph, "The code is very maintainable.", (0, 128, 0))  # Green
        elif file_metrics["Maintainability Index"] >= 50:
            color_text(summary_paragraph, "The code has moderate maintainability.", (255, 165, 0))  # Orange
        else:
            color_text(summary_paragraph, "The code has low maintainability and may need refactoring.", (255, 0, 0))  # Red
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        for metric, value in file_metrics.items():
            if metric != "file_path":
                row_cells = table.add_row().cells
                row_cells[0].text = str(metric)
                row_cells[1].text = f"{value:.2f}" if isinstance(value, float) else str(value)
    
    doc.save(output_path)
    print(f"Word report generated: {output_path}")

def generate_html_report(all_file_metrics, aggregated_metrics, output_path):
    if not aggregated_metrics:
        print("No data to generate an HTML report.")
        return
    
    html_content = """
    <html>
    <head>
        <title>Code Analysis Report</title>
        <style>
            body { font-family: Arial, sans-serif; margin: 20px; background-color: #f9f9f9; }
            .container { max-width: 1200px; margin: auto; background-color: #fff; padding: 20px; box-shadow: 0px 0px 10px rgba(0, 0, 0, 0.1); }
            h1, h2, h3 { color: #333; }
            table { width: 100%; border-collapse: collapse; margin-bottom: 20px; }
            th, td { padding: 10px; border: 1px solid #dddddd; text-align: left; }
            th { background-color: #f2f2f2; font-weight: bold; }
            .good { background-color: #d4edda; color: #155724; }
            .moderate { background-color: #fff3cd; color: #856404; }
            .poor { background-color: #f8d7da; color: #721c24; }
            .explanation { background-color: #e9ecef; padding: 10px; margin-bottom: 20px; border-left: 4px solid #007bff; }
            .summary { padding: 5px; font-weight: bold; }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>Code Analysis Report</h1>
            <p>Generated on: {date}</p>
    """.format(date=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    
    # Insert the explanations, aggregated metrics, and individual file metrics as before

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    print(f"HTML report generated: {output_path}")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Generate code analysis reports in HTML or Word format")
    parser.add_argument('folder_path', type=str, help="Path to the folder containing Python files to analyze")
    parser.add_argument('--format', choices=['html', 'word'], default='html', help="The report format to generate (html or word)")
    parser.add_argument('--output', type=str, required=True, help="The output path for the report")
    args = parser.parse_args()
    
    folder_path = args.folder_path
    report_format = args.format
    output_path = args.output

    all_file_metrics = analyze_project_folder(folder_path, exclude_folders=['venv'])
    aggregated_metrics = aggregate_metrics(all_file_metrics)

    if report_format == 'html':
        generate_html_report(all_file_metrics, aggregated_metrics, output_path)
    else:
        generate_word_report(all_file_metrics, aggregated_metrics, output_path)
